import os
import customtkinter as ctk
from tkinter import messagebox, filedialog
from docx import Document
from config import DATABASE_DIR

class DataExportMixin:
    def filtreli_veriyi_kaydet_word(self):
        if self.df_filtered is None or not self.secili_klasor is not self.secili_dosya :
            messagebox.showwarning("Uyarı", "Kaydedilecek filtreli veri yok.")
            return

        ana_isim = self.secili_klasor
        dosya_adi_base = self.secili_dosya.replace('.xlsx', '_filtrede')
        klasor_yolu = os.path.join(DATABASE_DIR, ana_isim)
        sayac = 1
        uzanti = '.docx'
        dosya_adi = dosya_adi_base + uzanti
        kayit_yolu = os.path.join(klasor_yolu, dosya_adi)

        while os.path.exists(kayit_yolu):
            dosya_adi = f"{dosya_adi_base}_{sayac}{uzanti}"
            kayit_yolu = os.path.join(klasor_yolu, dosya_adi)
            sayac += 1

        filtre_aciklamalari = []
        for f in self.filtreler:
            col, op, val1, val2, tip = f
            if op == 'Aralık':
                filtre_aciklamalari.append(f"Filtre: {col} aralığı: {val1} - {val2}")
            else:
                filtre_aciklamalari.append(f"Filtre: {col} anahtar kelime: {val1}")

        try:
            doc = Document()
            doc.add_heading("Filtrelenmiş Dosya", 0)
            for aciklama in filtre_aciklamalari:
                doc.add_paragraph(aciklama)

            if not self.df_filtered.empty:
                table = doc.add_table(rows=1, cols=len(self.df_filtered.columns))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for i, column_name in enumerate(self.df_filtered.columns):
                    hdr_cells[i].text = str(column_name)

                for _, row in self.df_filtered.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)
            else:
                doc.add_paragraph("Filtre sonrası veri bulunamadı.")

            doc.save(kayit_yolu)
            messagebox.showinfo("Kayıt", f"Filtreli veri Word'e kaydedildi: {dosya_adi}")
            self._dosya_agacini_guncelle()
        except Exception as e:
            messagebox.showerror("Hata", f"Word'e kayıt başarısız: {e}")

    def filtreli_veriyi_kaydet_excel(self):
        if self.df_filtered is None or self.secili_klasor is None or self.secili_dosya is None:
            messagebox.showwarning("Uyarı", "Kaydedilecek filtreli veri yok.")
            return

        ana_isim = self.secili_klasor
        dosya_adi_base = self.secili_dosya.replace('.xlsx', '_filtrede')
        klasor_yolu = os.path.join(DATABASE_DIR, ana_isim)
        sayac = 1
        uzanti = '.xlsx'
        dosya_adi = dosya_adi_base + uzanti
        kayit_yolu = os.path.join(klasor_yolu, dosya_adi)

        while os.path.exists(kayit_yolu):
            dosya_adi = f"{dosya_adi_base}_{sayac}{uzanti}"
            kayit_yolu = os.path.join(klasor_yolu, dosya_adi)
            sayac += 1

        try:
            self.df_filtered.to_excel(kayit_yolu, index=False)
            messagebox.showinfo("Kayıt", f"Filtreli veri Excel'e kaydedildi: {dosya_adi}")
            self._dosya_agacini_guncelle()
        except Exception as e:
            messagebox.showerror("Hata", f"Excel'e kayıt başarısız: {e}")

    def veri_kaydet_pencere(self):
        """Veriyi kaydetmek için format seçim penceresi açar"""
        if self.df is None:
            messagebox.showwarning("Uyarı", "Kaydedilecek veri yok.")
            return

        # Hangi veriyi kaydedeceğimizi belirle (filtrelenmiş veya orijinal)
        kaydedilecek_df = self.df_filtered if self.df_filtered is not None else self.df
        
        if kaydedilecek_df is None or kaydedilecek_df.empty:
            messagebox.showwarning("Uyarı", "Kaydedilecek veri yok.")
            return


        pencere = ctk.CTkToplevel(self)
        pencere.title("Veriyi Kaydet")
        pencere.geometry("450x400")
        pencere.resizable(False, False)
        pencere.grab_set()  # Modal pencere yap
        
        # Ana frame
        ana_frame = ctk.CTkFrame(pencere)
        ana_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Başlık
        ctk.CTkLabel(ana_frame, text="Kaydetme Formatı Seçin", 
                    font=ctk.CTkFont(size=16, weight="bold")).pack(pady=(20, 30))
        
        # Format seçimi
        format_frame = ctk.CTkFrame(ana_frame)
        format_frame.pack(fill="x", padx=20, pady=(10, 20))
        
        ctk.CTkLabel(format_frame, text="Dosya Formatı:", 
                    font=ctk.CTkFont(size=14, weight="bold")).pack(pady=(15, 10))
        
        format_var = ctk.StringVar(value="excel")
        
        excel_radio = ctk.CTkRadioButton(format_frame, text="Excel (.xlsx)", 
                                       variable=format_var, value="excel", height=25)
        excel_radio.pack(pady=8)
        
        word_radio = ctk.CTkRadioButton(format_frame, text="Word (.docx)", 
                                      variable=format_var, value="word", height=25)
        word_radio.pack(pady=8)
        
        csv_radio = ctk.CTkRadioButton(format_frame, text="CSV (.csv)", 
                                     variable=format_var, value="csv", height=25)
        csv_radio.pack(pady=8)
        
        # Butonlar
        buton_frame = ctk.CTkFrame(ana_frame)
        buton_frame.pack(fill="x", padx=20, pady=(20, 10))
        
        def kaydet():
            secilen_format = format_var.get()
            
            # Dosya adı önerisi
            if self.secili_dosya:
                base_name = os.path.splitext(self.secili_dosya)[0]
                if self.df_filtered is not None:
                    base_name += "_filtrede"
            else:
                base_name = "veri"
            
            # Format uzantıları **
            format_extensions = {
                "excel": ".xlsx",
                "word": ".docx", 
                "csv": ".csv"
            }
            
            try:
                if secilen_format == "excel":
                    # Excel için direkt dosya kaydetme dialogu
                    dosya_yolu = filedialog.asksaveasfilename(
                        title="Excel Dosyasını Kaydet",
                        defaultextension=".xlsx",
                        filetypes=[("Excel Dosyası", "*.xlsx")],
                        initialfile=base_name + ".xlsx"
                    )
                    
                    if not dosya_yolu:
                        return
                    
                    kaydedilecek_df.to_excel(dosya_yolu, index=False)
                    messagebox.showinfo("Başarılı", f"Veri Excel formatında kaydedildi:\n{dosya_yolu}")
                    
                elif secilen_format == "word":
                    # Word için önce sütun seçim penceresi
                    secilen_sutunlar = self._word_sutun_secim_pencere(kaydedilecek_df.columns)
                    if not secilen_sutunlar:  # Kullanıcı iptali
                        return
                    
                    # Seçilen sütunlarla veriyi filtrele
                    kaydedilecek_df_word = kaydedilecek_df[secilen_sutunlar]
                    
                    # Sonra dosya kaydetme dialogu
                    dosya_yolu_word = filedialog.asksaveasfilename(
                        title="Word Dosyasını Kaydet",
                        defaultextension=".docx",
                        filetypes=[("Word Dosyası", "*.docx")],
                        initialfile=base_name + ".docx"
                    )
                    
                    if not dosya_yolu_word:
                        return
                    
                    # Word formatında kaydetme
                    doc = Document()
                    doc.add_heading("Veri Çıktısı", 0)
                    
                    # Filtre açıklamaları ekle
                    if hasattr(self, 'filtreler') and self.filtreler:
                        doc.add_heading("Uygulanan Filtreler:", level=1)
                        for f in self.filtreler:
                            col, op, val1, val2, tip = f
                            if op == 'Aralık':
                                doc.add_paragraph(f"• {col} aralığı: {val1} - {val2}")
                            else:
                                doc.add_paragraph(f"• {col} anahtar kelime: {val1}")
                        doc.add_paragraph("")  # Boş satır
                    
                    # Tablo oluştur
                    if not kaydedilecek_df_word.empty:
                        table = doc.add_table(rows=1, cols=len(kaydedilecek_df_word.columns))
                        table.style = 'Table Grid'
                        hdr_cells = table.rows[0].cells
                        for i, column_name in enumerate(kaydedilecek_df_word.columns):
                            hdr_cells[i].text = str(column_name)

                        for _, row in kaydedilecek_df_word.iterrows():
                            row_cells = table.add_row().cells
                            for i, value in enumerate(row):
                                row_cells[i].text = str(value)
                    else:
                        doc.add_paragraph("Veri bulunamadı.")
                    
                    doc.save(dosya_yolu_word)
                    messagebox.showinfo("Başarılı", f"Veri Word formatında kaydedildi:\n{dosya_yolu_word}")
                    
                elif secilen_format == "csv":
                    # CSV için direkt dosya kaydetme dialog'u
                    dosya_yolu = filedialog.asksaveasfilename(
                        title="CSV Dosyasını Kaydet",
                        defaultextension=".csv",
                        filetypes=[("CSV Dosyası", "*.csv")],
                        initialfile=base_name + ".csv"
                    )
                    
                    if not dosya_yolu:
                        return
                    
                    kaydedilecek_df.to_csv(dosya_yolu, index=False, encoding='utf-8-sig')
                    messagebox.showinfo("Başarılı", f"Veri CSV formatında kaydedildi:\n{dosya_yolu}")
                
                # Dosya ağacını güncelle (tüm formatlar için)
                self._dosya_agacini_guncelle()
                    
            except Exception as e:
                messagebox.showerror("Hata", f"Kayıt başarısız:\n{e}")
            
            pencere.destroy()
        
        def iptal():
            pencere.destroy()
        
        kaydet_btn = ctk.CTkButton(buton_frame, text="Kaydet", command=kaydet, height=45,
                                    font=ctk.CTkFont(size=14, weight="bold"), 
                                    fg_color="#2E8B57", hover_color="#228B22")
        kaydet_btn.pack(side="left", padx=5, pady=10, fill="x", expand=True)
        
        iptal_btn = ctk.CTkButton(buton_frame, text="İptal", command=iptal, 
                                  height=45, font=ctk.CTkFont(size=14, weight="bold"),
                                    fg_color="#DC143C", hover_color="#B22222")
        
        iptal_btn.pack(side="right", padx=5, pady=10, fill="x", expand=True)

    def _word_sutun_secim_pencere(self, sutunlar):
        """Word formatında kaydetmek için sütun seçim penceresi"""

        pencere = ctk.CTkToplevel(self)
        pencere.title("Sütun Seçimi")
        pencere.geometry("400x550")
        pencere.resizable(False, False)
        pencere.grab_set()  
        
        # Ana frame
        ana_frame = ctk.CTkFrame(pencere)
        ana_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Başlık
        ctk.CTkLabel(ana_frame, text="Word'e Kaydedilecek Sütunları Seçin", 
                    font=ctk.CTkFont(size=16, weight="bold")).pack(pady=(20, 10))
        
        # Açıklama
        ctk.CTkLabel(ana_frame, text="İstediğiniz sütunları seçin:", 
                    font=ctk.CTkFont(size=12)).pack(pady=(0, 20))
        
        # Scrollable frame için container
        scroll_container = ctk.CTkFrame(ana_frame)
        scroll_container.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Scrollable frame
        scroll_frame = ctk.CTkScrollableFrame(scroll_container, height=250)
        scroll_frame.pack(fill="both", expand=True)
        
        # Sütun seçim değişkenleri
        sutun_vars = {}
        secilen_sutunlar = []
        
        # Tüm sütunları varsayılan olarak seçili yap
        for sutun in sutunlar:
            var = ctk.BooleanVar(value=True)  # Varsayılan olarak seçili
            sutun_vars[sutun] = var
            
            checkbox = ctk.CTkCheckBox(scroll_frame, text=sutun, variable=var, height=25)
            checkbox.pack(anchor="w", padx=10, pady=5, fill="x")
        
        # Butonlar
        buton_frame = ctk.CTkFrame(ana_frame)
        buton_frame.pack(fill="x", padx=10, pady=(20, 10))
        
        def tumunu_sec():
            for var in sutun_vars.values():
                var.set(True)
        
        def tumunu_kaldir():
            for var in sutun_vars.values():
                var.set(False)
        
        def kaydet():
            nonlocal secilen_sutunlar
            secilen_sutunlar = [sutun for sutun, var in sutun_vars.items() if var.get()]
            if not secilen_sutunlar:
                messagebox.showwarning("Uyarı", "En az bir sütun seçmelisiniz!")
                return
            pencere.destroy()
        
        def iptal():
            nonlocal secilen_sutunlar
            secilen_sutunlar = []
            pencere.destroy()
        
        # Butonlar
        ctk.CTkButton(buton_frame, text="Tümünü Seç", command=tumunu_sec,
                       height=35, font=ctk.CTkFont(size=12, weight="bold"),
                         fg_color="#4682B4", hover_color="#5F9EA0").pack(side="left", padx=5, pady=5, fill="x", expand=True)
        ctk.CTkButton(buton_frame, text="Tümünü Kaldır", command=tumunu_kaldir,
                       height=35, font=ctk.CTkFont(size=12, weight="bold"),
                         fg_color="#696969", hover_color="#808080").pack(side="left", padx=5, pady=5, fill="x", expand=True)
        
        buton_frame2 = ctk.CTkFrame(ana_frame)
        buton_frame2.pack(fill="x", padx=10, pady=(0, 1))
        
        ctk.CTkButton(buton_frame2, text="Kaydet", command=kaydet, height=80,
                       font=ctk.CTkFont(size=16, weight="bold"),
                         fg_color="#2E8B57", hover_color="#228B22").pack(side="left", padx=5, pady=10, fill="x", expand=True)
        
        ctk.CTkButton(buton_frame2, text="İptal", command=iptal, 
                      height=80, font=ctk.CTkFont(size=16, weight="bold"), 
                      fg_color="#DC143C", hover_color="#B22222").pack(side="right", padx=5, pady=10, fill="x", expand=True)
        
        # Pencere kapatılana kadar bekle
        pencere.wait_window()
        
        return secilen_sutunlar